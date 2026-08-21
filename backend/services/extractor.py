"""
TalonAI — Document Extraction Service
======================================
Extraction priority per file type:

  IMAGE FILES  (png/jpg/jpeg/webp/bmp/tiff)
  ─────────────────────────────────────────
  1. LLMWhisperer API  (requires LLMWHISPERER_API_KEY)
     • Best accuracy, layout-preserving, handles complex scans
     • Same cloud API used for PDFs — accepts image files directly
     • 100 pages/month free tier
  2. Tesseract 5        — local fallback, no API key needed
  3. Metadata stub      — last-resort if both engines missing

  NOTE: PP-OCRv6 tiny is installed (paddleocr 3.7) but the
  safetensors engine binding is not yet registered in this version.
  Will be re-enabled once paddleocr ships the binding fix.
  The MULTILINGUAL NOTE below still applies when it is re-enabled.

  PDF FILES
  ─────────────────────────────────────────
  1. LLMWhisperer API  (requires LLMWHISPERER_API_KEY env var)
     • Layout-preserving, best for tables / multi-column / scanned PDFs
  2. pdfplumber        — text-native PDFs, table extraction
  3. pypdf             — plain text fallback
  4. Metadata stub     — scanned PDF with no text layer

  DOCX, XLSX/CSV, TXT/MD — unchanged

──────────────────────────────────────────────────────────────────────
MULTILINGUAL NOTE (PP-OCRv6 tiny — available once binding is fixed)
──────────────────────────────────────────────────────────────────────
PP-OCRv6 tiny supports 49 languages including:
  Chinese (Simplified + Traditional), Japanese, Korean, Arabic, Hindi,
  German, French, Spanish, Russian, and 40 other Latin-script languages.

To enable multilingual OCR once the engine binding is fixed,
change `_PADDLEOCR_LANG` below from "en" to:
  • "ch"           — Chinese + English (mixed docs)
  • "multilingual" — all 49 languages auto-detected
  • "japan", "arabic", "french", etc. — single language
──────────────────────────────────────────────────────────────────────
"""

import os
import io
import re
import shutil
import tempfile
import logging
from typing import Dict, Any, List, Tuple, Optional

from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from pypdf import PdfReader
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd

from models import DocumentMetadata, ExtractedDocument

logger = logging.getLogger(__name__)

# ══════════════════════════════════════════════════════════════════════
#  Language setting for PP-OCRv6 tiny
#  Change to "multilingual" or a specific language code to enable
#  multi-language OCR. See MULTILINGUAL NOTE in module docstring.
# ══════════════════════════════════════════════════════════════════════
_PADDLEOCR_LANG = "en"


# ══════════════════════════════════════════════════════════════════════
#  Engine detection — runs once at import time, no side-effects
# ══════════════════════════════════════════════════════════════════════

def _detect_tesseract() -> bool:
    """Auto-detect Tesseract OCR binary on Windows and Linux."""
    if shutil.which("tesseract"):
        return True
    windows_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
        os.path.expandvars(r"%USERPROFILE%\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
    ]
    for path in windows_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            return True
    return False


def _detect_paddleocr() -> bool:
    """Check paddleocr 3.x is importable (does NOT load model weights yet)."""
    try:
        from paddleocr import TextDetection  # noqa: F401
        return True
    except ImportError:
        return False


def _detect_llmwhisperer() -> bool:
    """Check LLMWhisperer client is importable and API key is set."""
    try:
        from unstract.llmwhisperer import LLMWhispererClientV2  # noqa: F401
        return bool(os.environ.get("LLMWHISPERER_API_KEY", "").strip())
    except ImportError:
        return False


HAS_TESSERACT    = _detect_tesseract()
HAS_PADDLEOCR   = _detect_paddleocr()
HAS_LLMWHISPERER = _detect_llmwhisperer()

logger.info(
    "Extraction engines: Tesseract=%s | PP-OCRv6-tiny=%s | LLMWhisperer=%s",
    HAS_TESSERACT, HAS_PADDLEOCR, HAS_LLMWHISPERER,
)


# ══════════════════════════════════════════════════════════════════════
#  Shared helpers
# ══════════════════════════════════════════════════════════════════════

def count_words_and_sentences(text: str) -> Tuple[int, int, int]:
    char_count  = len(text)
    words       = re.findall(r'\b\w+\b', text)
    word_count  = len(words)
    sentences   = [s for s in re.split(r'[.!?]+', text) if s.strip()]
    sentence_count = max(1, len(sentences))
    return char_count, word_count, sentence_count


def calculate_reading_time(word_count: int) -> float:
    return round(word_count / 200.0, 1)


def _build_doc(
    filename:   str,
    file_type:  str,
    file_bytes: bytes,
    raw_text:   str,
    sections:   list,
    extraction_method: str,
    page_or_sheet_count: int = 1,
) -> ExtractedDocument:
    """Build a normalised ExtractedDocument from any extraction result."""
    char_c, word_c, sent_c = count_words_and_sentences(raw_text)
    metadata = DocumentMetadata(
        filename=filename,
        file_type=file_type,
        file_size_bytes=len(file_bytes),
        char_count=char_c,
        word_count=word_c,
        sentence_count=sent_c,
        estimated_read_time_minutes=calculate_reading_time(word_c),
        page_or_sheet_count=max(1, page_or_sheet_count),
        extraction_method=extraction_method,
    )
    return ExtractedDocument(
        metadata=metadata,
        raw_text=raw_text,
        sections=sections or [{"title": "Document Content", "content": raw_text, "page": 1}],
        preview=raw_text[:500] + ("..." if len(raw_text) > 500 else ""),
    )


# ══════════════════════════════════════════════════════════════════════
#  IMAGE EXTRACTION
# ══════════════════════════════════════════════════════════════════════

def _preprocess_image(image: Image.Image) -> Image.Image:
    """Grayscale → contrast boost → sharpen — improves accuracy for both engines."""
    try:
        gray     = image.convert("L")
        enhanced = ImageEnhance.Contrast(gray).enhance(2.0)
        return enhanced.filter(ImageFilter.SHARPEN)
    except Exception:
        return image


# ── PP-OCRv6 tiny via Transformers backend ─────────────────────────────

_paddle_det_model = None  # lazy-loaded singleton
_paddle_rec_model = None


def _load_paddle_models():
    """Load PP-OCRv6 tiny detection + recognition models (once, on first call)."""
    global _paddle_det_model, _paddle_rec_model
    if _paddle_det_model is not None:
        return True
    try:
        from paddleocr import TextDetection, TextRecognition
        # PP-OCRv6 tiny — 0.43 M det params + 1.1 M rec params
        # backend="transformers" avoids the heavy paddlepaddle runtime
        _paddle_det_model = TextDetection(
            model_name="PP-OCRv6_tiny_det_safetensors",
            backend="transformers",
        )
        _paddle_rec_model = TextRecognition(
            model_name="PP-OCRv6_tiny_rec_safetensors",
            backend="transformers",
            lang=_PADDLEOCR_LANG,
        )
        logger.info("PP-OCRv6 tiny models loaded (lang=%s)", _PADDLEOCR_LANG)
        return True
    except Exception as e:
        logger.warning("PP-OCRv6 model load failed: %s", e)
        _paddle_det_model = None
        _paddle_rec_model = None
        return False


def _extract_image_paddleocr(file_bytes: bytes, filename: str) -> Optional[ExtractedDocument]:
    """
    Run PP-OCRv6 tiny on an image file.
    Detection → bounding boxes, Recognition → text per box,
    assembled top-to-bottom by y-coordinate.
    Returns None on any failure so caller can fall back to Tesseract.
    """
    try:
        import numpy as np

        if not _load_paddle_models():
            return None

        image     = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        img_array = np.array(image)

        # Detection: returns list of bounding box results
        det_output = _paddle_det_model.predict(img_array, batch_size=1)
        if not det_output:
            return None

        # Collect (y_top, text) pairs
        lines: List[Tuple[float, str]] = []
        for det_res in det_output:
            boxes = det_res.get("dt_polys", [])
            if not boxes:
                continue
            for box in boxes:
                # Crop the detected region and run recognition
                pts   = [list(map(int, p)) for p in box]
                xs    = [p[0] for p in pts]
                ys    = [p[1] for p in pts]
                x1, y1 = max(0, min(xs)), max(0, min(ys))
                x2, y2 = min(image.width, max(xs)), min(image.height, max(ys))
                if x2 <= x1 or y2 <= y1:
                    continue
                region    = np.array(image.crop((x1, y1, x2, y2)))
                rec_out   = _paddle_rec_model.predict(region, batch_size=1)
                for rec_res in rec_out:
                    text = rec_res.get("rec_text", "").strip()
                    if text:
                        lines.append((float(y1), text))

        # Sort top-to-bottom
        lines.sort(key=lambda t: t[0])
        raw_text = "\n".join(t for _, t in lines).strip()

        if not raw_text:
            return None

        return _build_doc(
            filename=filename,
            file_type="Image / Scanned Document",
            file_bytes=file_bytes,
            raw_text=raw_text,
            sections=[{"title": "Image OCR Text", "content": raw_text, "page": 1}],
            extraction_method=f"PP-OCRv6 tiny (lang={_PADDLEOCR_LANG})",
        )

    except Exception as e:
        logger.warning("PP-OCRv6 extraction failed (%s), falling back to Tesseract: %s", filename, e)
        return None


# ── Tesseract fallback ─────────────────────────────────────────────────

def _extract_image_tesseract(file_bytes: bytes, filename: str) -> ExtractedDocument:
    image     = Image.open(io.BytesIO(file_bytes))
    processed = _preprocess_image(image)
    raw_text  = ""
    method    = "Tesseract OCR"

    try:
        raw_text = pytesseract.image_to_string(processed).strip()
    except Exception as e:
        logger.warning("Tesseract failed: %s", e)
        raw_text = (
            f"[OCR Note: OCR attempted on {filename} "
            f"({image.width}x{image.height} px). "
            f"Install Tesseract or provide a LLMWHISPERER_API_KEY for cloud OCR.]"
        )
        method = "Image Metadata / OCR Fallback"

    if not raw_text:
        raw_text = f"[Image File: {filename} ({image.width}x{image.height} px). No readable text detected.]"

    return _build_doc(
        filename=filename,
        file_type="Image / Scanned Document",
        file_bytes=file_bytes,
        raw_text=raw_text,
        sections=[{"title": "Image OCR Text", "content": raw_text, "page": 1}],
        extraction_method=method,
    )


# ── LLMWhisperer for images ────────────────────────────────────────────

def _extract_image_llmwhisperer(file_bytes: bytes, filename: str) -> Optional[ExtractedDocument]:
    """
    Send image to LLMWhisperer cloud API.
    LLMWhisperer accepts PNG/JPG/WEBP/BMP/TIFF directly — same endpoint as PDFs.
    Returns None on missing key, error, or empty result so caller falls back.
    """
    api_key = os.environ.get("LLMWHISPERER_API_KEY", "").strip()
    if not api_key:
        return None

    ext = os.path.splitext(filename)[1].lower() or ".png"
    tmp_path = None
    try:
        from unstract.llmwhisperer import LLMWhispererClientV2

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        client = LLMWhispererClientV2(
            base_url="https://llmwhisperer-api.us-central.unstract.com/api/v2",
            api_key=api_key,
        )
        result = client.whisper(
            file_path=tmp_path,
            wait_for_completion=True,
            wait_timeout=120,
        )
        raw_text = result.get("extraction", {}).get("result_text", "").strip()
        if not raw_text:
            return None

        logger.info("LLMWhisperer extracted image: %s (%d chars)", filename, len(raw_text))

        return _build_doc(
            filename=filename,
            file_type="Image / Scanned Document",
            file_bytes=file_bytes,
            raw_text=raw_text,
            sections=[{"title": "Image OCR Text", "content": raw_text, "page": 1}],
            extraction_method="LLMWhisperer OCR (layout-preserving)",
        )

    except Exception as e:
        logger.warning("LLMWhisperer image OCR failed (%s): %s — falling back to Tesseract", filename, e)
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ── Public image entry point ───────────────────────────────────────────

def extract_from_image(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """
    Priority:
    1. LLMWhisperer  (if LLMWHISPERER_API_KEY set) — best accuracy, layout-aware
    2. PP-OCRv6 tiny (if installed)                — accurate local layout OCR
    3. Tesseract 5   (if installed locally)         — fast local fallback
    4. Metadata stub (if none available)
    """
    # 1 — LLMWhisperer (best quality, handles complex scans)
    if HAS_LLMWHISPERER:
        result = _extract_image_llmwhisperer(file_bytes, filename)
        if result is not None:
            return result

    # 2 — PP-OCRv6 tiny (local neural OCR)
    if HAS_PADDLEOCR:
        result = _extract_image_paddleocr(file_bytes, filename)
        if result is not None:
            return result

    # 3 — Tesseract local fallback
    return _extract_image_tesseract(file_bytes, filename)


# ══════════════════════════════════════════════════════════════════════
#  PDF EXTRACTION
# ══════════════════════════════════════════════════════════════════════

# ── LLMWhisperer (layout-preserving, best for tables) ─────────────────

def _extract_pdf_llmwhisperer(file_bytes: bytes, filename: str) -> Optional[ExtractedDocument]:
    """
    Send the PDF to LLMWhisperer's cloud API.
    Returns None on missing key, quota error, or any exception so the
    caller falls through to pdfplumber.
    """
    api_key = os.environ.get("LLMWHISPERER_API_KEY", "").strip()
    if not api_key:
        return None

    tmp_path = None
    try:
        from unstract.llmwhisperer import LLMWhispererClientV2

        # Write bytes to a temp file — the API requires a file path
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        client = LLMWhispererClientV2(
            base_url="https://llmwhisperer-api.us-central.unstract.com/api/v2",
            api_key=api_key,
        )

        result = client.whisper(
            file_path=tmp_path,
            wait_for_completion=True,
            wait_timeout=120,
        )

        raw_text = result.get("extraction", {}).get("result_text", "").strip()
        if not raw_text:
            return None

        # LLMWhisperer returns a single layout-preserved text block.
        # Split into page sections using the page-break markers it inserts.
        page_chunks  = re.split(r'\f|\[PAGE_BREAK\]', raw_text)
        sections     = [
            {"title": f"Page {i+1}", "content": chunk.strip(), "page": i+1}
            for i, chunk in enumerate(page_chunks) if chunk.strip()
        ]
        page_count = len(sections) or 1

        logger.info("LLMWhisperer extracted %d pages from %s", page_count, filename)

        return _build_doc(
            filename=filename,
            file_type="PDF Document",
            file_bytes=file_bytes,
            raw_text=raw_text,
            sections=sections,
            extraction_method="LLMWhisperer (layout-preserving)",
            page_or_sheet_count=page_count,
        )

    except Exception as e:
        logger.warning("LLMWhisperer failed for %s: %s — falling back to pdfplumber", filename, e)
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ── pdfplumber (text-native PDFs) ─────────────────────────────────────

def _extract_pdf_pdfplumber(file_bytes: bytes, filename: str) -> Tuple[List, List, int]:
    """Returns (sections, full_text_list, page_count). Empty lists on failure."""
    sections       = []
    full_text_list = []
    page_count     = 0
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for idx, page in enumerate(pdf.pages, start=1):
                p_text = page.extract_text(layout=True) or ""
                tables = page.extract_tables()
                if tables:
                    table_lines = []
                    for t in tables:
                        rows = [" | ".join(str(c or "").strip() for c in row) for row in t if any(row)]
                        if rows:
                            table_lines.append("\n" + "\n".join(rows) + "\n")
                    if table_lines:
                        p_text += "\n\n[Extracted Tables]:\n" + "\n".join(table_lines)
                p_text = p_text.strip()
                if p_text:
                    sections.append({"title": f"Page {idx}", "content": p_text, "page": idx})
                    full_text_list.append(f"--- Page {idx} ---\n{p_text}")
    except Exception as e:
        logger.debug("pdfplumber failed: %s", e)
    return sections, full_text_list, page_count


def _extract_pdf_pypdf(file_bytes: bytes) -> Tuple[List, List, int]:
    """pypdf fallback. Returns (sections, full_text_list, page_count)."""
    sections       = []
    full_text_list = []
    page_count     = 0
    try:
        reader     = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        for idx, page in enumerate(reader.pages, start=1):
            p_text = (page.extract_text() or "").strip()
            if p_text:
                sections.append({"title": f"Page {idx}", "content": p_text, "page": idx})
                full_text_list.append(f"--- Page {idx} ---\n{p_text}")
    except Exception as e:
        logger.debug("pypdf failed: %s", e)
    return sections, full_text_list, page_count


# ── Public PDF entry point ─────────────────────────────────────────────

def extract_from_pdf(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """
    Priority:
    1. LLMWhisperer   (if LLMWHISPERER_API_KEY set)
    2. pdfplumber     (text-native PDFs, table extraction)
    3. pypdf          (plain text fallback)
    4. Metadata stub  (scanned image-only PDF)
    """
    # 1 — LLMWhisperer
    if HAS_LLMWHISPERER:
        result = _extract_pdf_llmwhisperer(file_bytes, filename)
        if result is not None:
            return result

    # 2 — pdfplumber
    sections, full_text_list, page_count = _extract_pdf_pdfplumber(file_bytes, filename)

    # 3 — pypdf fallback if pdfplumber got nothing
    if not full_text_list:
        sections, full_text_list, page_count = _extract_pdf_pypdf(file_bytes)

    raw_text = "\n\n".join(full_text_list).strip()

    if not raw_text:
        raw_text = (
            f"[Scanned PDF: {filename} ({page_count} pages). "
            f"No embedded text layer detected. "
            f"Set LLMWHISPERER_API_KEY for cloud OCR of scanned PDFs.]"
        )
        method = "PDF Structure (Scanned / Image-based)"
    else:
        method = "PDF Parser (pdfplumber/PyPDF)"

    return _build_doc(
        filename=filename,
        file_type="PDF Document",
        file_bytes=file_bytes,
        raw_text=raw_text,
        sections=sections,
        extraction_method=method,
        page_or_sheet_count=max(1, page_count),
    )


# ══════════════════════════════════════════════════════════════════════
#  DOCX EXTRACTION  — unchanged
# ══════════════════════════════════════════════════════════════════════

def extract_from_docx(file_bytes: bytes, filename: str) -> ExtractedDocument:
    doc          = DocxDocument(io.BytesIO(file_bytes))
    sections     = []
    full_text    = []
    curr_heading = "Document Beginning"
    curr_content: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            if curr_content:
                sections.append({
                    "title": curr_heading,
                    "content": "\n".join(curr_content),
                    "page": len(sections) + 1,
                })
                curr_content = []
            curr_heading = text
            full_text.append(f"\n## {text}\n")
        else:
            curr_content.append(text)
            full_text.append(text)

    if curr_content:
        sections.append({
            "title": curr_heading,
            "content": "\n".join(curr_content),
            "page": len(sections) + 1,
        })

    for table_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tbl = f"\n[Table {table_idx}]:\n" + "\n".join(rows)
            full_text.append(tbl)
            sections.append({"title": f"Table {table_idx}", "content": tbl, "page": len(sections) + 1})

    raw_text = "\n".join(full_text).strip()
    if not raw_text:
        raw_text = f"[Word Document: {filename} is empty or contains unsupported content.]"

    return _build_doc(
        filename=filename,
        file_type="Microsoft Word (.docx)",
        file_bytes=file_bytes,
        raw_text=raw_text,
        sections=sections,
        extraction_method="python-docx Parser",
        page_or_sheet_count=max(1, len(sections)),
    )


# ══════════════════════════════════════════════════════════════════════
#  EXCEL / CSV EXTRACTION  — unchanged
# ══════════════════════════════════════════════════════════════════════

def extract_from_excel(file_bytes: bytes, filename: str) -> ExtractedDocument:
    sections       = []
    full_text_list = []
    sheet_count    = 1

    try:
        if filename.lower().endswith(".csv"):
            df         = pd.read_csv(io.BytesIO(file_bytes))
            sheet_dict = {"CSV Data": df}
        else:
            sheet_dict  = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            sheet_count = len(sheet_dict)

        for sheet_name, df in sheet_dict.items():
            header   = f"=== Sheet: {sheet_name} (Rows: {len(df)}, Columns: {len(df.columns)}) ==="
            cols     = ", ".join(str(c) for c in df.columns)
            summary  = f"Columns ({len(df.columns)}): {cols}\n"
            num_cols = df.select_dtypes(include=["number"])
            stats    = ("\n[Key Numeric Statistics]:\n" + num_cols.describe().round(2).to_string()) if not num_cols.empty else ""
            preview  = df.head(15).to_string(index=False)
            body     = f"{summary}\n[Sample Data Preview]:\n{preview}\n{stats}"
            full_text_list.append(header)
            full_text_list.append(body)
            sections.append({"title": f"Sheet: {sheet_name}", "content": f"{header}\n\n{body}", "page": len(sections) + 1})
    except Exception as e:
        full_text_list.append(f"[Error parsing spreadsheet {filename}: {e}]")

    raw_text = "\n\n".join(full_text_list).strip()

    return _build_doc(
        filename=filename,
        file_type="Spreadsheet (Excel/CSV)",
        file_bytes=file_bytes,
        raw_text=raw_text,
        sections=sections,
        extraction_method="Pandas & OpenPyXL Engine",
        page_or_sheet_count=sheet_count,
    )


# ══════════════════════════════════════════════════════════════════════
#  PLAIN TEXT / MARKDOWN EXTRACTION  — unchanged
# ══════════════════════════════════════════════════════════════════════

def extract_from_plain_text(file_bytes: bytes, filename: str) -> ExtractedDocument:
    raw_text = ""
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            raw_text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if not raw_text:
        raw_text = file_bytes.decode("utf-8", errors="replace")

    raw_text  = raw_text.strip()
    sections  = []
    lines     = raw_text.splitlines()
    curr_title: str       = "Introduction"
    curr_lines: List[str] = []

    for line in lines:
        if line.startswith("#"):
            if curr_lines:
                sections.append({"title": curr_title, "content": "\n".join(curr_lines).strip(), "page": len(sections) + 1})
                curr_lines = []
            curr_title = line.lstrip("#").strip() or f"Section {len(sections)+1}"
        else:
            curr_lines.append(line)

    if curr_lines:
        sections.append({"title": curr_title, "content": "\n".join(curr_lines).strip(), "page": len(sections) + 1})

    return _build_doc(
        filename=filename,
        file_type="Text / Markdown Document",
        file_bytes=file_bytes,
        raw_text=raw_text,
        sections=sections,
        extraction_method="Native UTF-8 Text Stream",
        page_or_sheet_count=max(1, len(sections)),
    )


# ══════════════════════════════════════════════════════════════════════
#  MASTER ROUTER
# ══════════════════════════════════════════════════════════════════════

def extract_document(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """Route to the appropriate extractor based on file extension."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_bytes, filename)
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        return extract_from_image(file_bytes, filename)
    elif ext in {".docx", ".doc"}:
        return extract_from_docx(file_bytes, filename)
    elif ext in {".xlsx", ".xls", ".csv", ".tsv"}:
        return extract_from_excel(file_bytes, filename)
    elif ext in {".txt", ".md", ".rtf", ".json", ".log", ".xml", ".html", ".py", ".js", ".css"}:
        return extract_from_plain_text(file_bytes, filename)
    else:
        try:
            return extract_from_plain_text(file_bytes, filename)
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")
