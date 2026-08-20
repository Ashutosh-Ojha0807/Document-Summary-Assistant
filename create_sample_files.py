import os
import io
import pandas as pd
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfWriter

os.makedirs("sample_docs", exist_ok=True)

# 1. Create DOCX
doc = Document()
doc.add_heading('Quarterly Strategic Business Plan 2026', 0)
doc.add_paragraph('Author: Strategic Planning Board')
doc.add_paragraph('Global tech operations showed strong performance across all cloud units in Q3.')

doc.add_heading('Key Strategic Initiatives', level=1)
doc.add_paragraph('• Expand regional infrastructure in APAC by December 2026.\n• Enhance security compliance for ISO 27001.\n• Automate document analysis pipelines.')

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Business Unit'
hdr_cells[1].text = 'Q3 Target ($M)'
hdr_cells[2].text = 'Actual Revenue ($M)'

data = [
    ('Cloud AI Services', '15.0', '18.4'),
    ('Enterprise Software', '12.0', '14.2'),
    ('Cybersecurity', '8.0', '9.9')
]
for bu, target, actual in data:
    row_cells = table.add_row().cells
    row_cells[0].text = bu
    row_cells[1].text = target
    row_cells[2].text = actual

doc.save("sample_docs/quarterly_strategic_plan.docx")
print("Created quarterly_strategic_plan.docx")

# 2. Create CSV / Excel
df = pd.DataFrame({
    "Region": ["North America", "Europe", "Asia-Pacific", "Latin America", "Middle East"],
    "Active_Users": [450000, 320000, 580000, 140000, 95000],
    "Revenue_USD": [14500000, 9800000, 18200000, 4200000, 3100000],
    "YoY_Growth_Pct": [24.5, 18.2, 36.8, 14.1, 28.3]
})
df.to_csv("sample_docs/regional_sales_q3.csv", index=False)
df.to_excel("sample_docs/regional_sales_q3.xlsx", index=False)
print("Created regional_sales_q3.csv and regional_sales_q3.xlsx")

# 3. Create Sample Image for OCR
img = Image.new('RGB', (800, 400), color=(255, 255, 255))
draw = ImageDraw.Draw(img)
# Draw text
sample_ocr_text = [
    "INVOICE & SERVICE STATEMENT #INV-2026-884",
    "Vendor: Apex Cloud Infrastructure Ltd.",
    "Date: August 18, 2026",
    "--------------------------------------------------",
    "Description                 Qty     Rate        Amount",
    "Dedicated GPU Cluster H100   4      $3,500.00   $14,000.00",
    "Cloud Storage 50TB           1        $850.00      $850.00",
    "Enterprise Support Tier      1      $1,200.00   $1,200.00",
    "--------------------------------------------------",
    "Total Payable: $16,050.00 (Due within 30 days)",
    "Action: Wire payment to Apex Treasury Account."
]

y = 30
for line in sample_ocr_text:
    draw.text((40, y), line, fill=(20, 20, 20))
    y += 30

img.save("sample_docs/sample_invoice_scan.png")
print("Created sample_invoice_scan.png")

# 4. Create Sample PDF
writer = PdfWriter()
page = writer.add_blank_page(width=595, height=842) # A4
# We can save a minimal PDF structure
with open("sample_docs/sample_document.pdf", "wb") as f:
    writer.write(f)
print("Created sample_document.pdf")
